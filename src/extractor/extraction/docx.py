"""Word (.docx) extraction."""

from __future__ import annotations

from pathlib import Path

from extractor.extraction.base import ExtractionError


def extract(path: Path) -> str:
    import docx

    try:
        document = docx.Document(str(path))
    except Exception as e:
        raise ExtractionError(f"python-docx could not open file: {e}") from e

    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    text = "\n".join(parts)
    if not text.strip():
        raise ExtractionError("No extractable text found in .docx")
    return text
