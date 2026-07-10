"""Generates a fictional company's Q3 review documents for demoing the extractor.

Run with the project's dev environment (needs fpdf2, python-docx, openpyxl,
all already dev dependencies): `python demo/generate_demo_files.py`

The story: Acme Robotics, Q3 2026. Revenue and pipeline data live in a PDF and
a spreadsheet; the two biggest risks (a supply-chain issue and rising customer
churn) are each corroborated across two *different* file formats; a hiring
freeze exception is buried in meeting notes; and one file is entirely
unrelated, to demonstrate that irrelevant documents get correctly excluded.
"""

from __future__ import annotations

from email.message import EmailMessage
from pathlib import Path

DOCS_DIR = Path(__file__).parent / "docs"


def make_financial_summary_pdf() -> None:
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, "Acme Robotics - Q3 2026 Financial Summary", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", size=11)
    pdf.ln(4)

    body = [
        "Total company revenue reached $18.4M this quarter, up 22% year over year.",
        "",
        "Revenue by division:",
        "  - Robotics Division: $12.1M (+31% YoY), driven by strong demand for the",
        "    Atlas arm product line.",
        "  - Services Division: $6.3M (+4% YoY), roughly in line with plan.",
        "",
        "Gross margin improved to 58%, up from 54% last quarter, mainly due to",
        "manufacturing efficiency gains at the Reno facility.",
        "",
        "Key financial risk: our proximity-sensor supplier, SensTech Inc., has",
        "signaled a 6-8 week increase in lead times starting in Q4. This could",
        "delay shipments of the Atlas arm if not mitigated.",
        "",
        "Operating expenses were $9.8M, up 11% YoY, primarily due to increased",
        "engineering headcount.",
    ]
    for line in body:
        pdf.cell(0, 6, line, new_x="LMARGIN", new_y="NEXT")

    pdf.output(str(DOCS_DIR / "financial_summary.pdf"))


def make_engineering_status_docx() -> None:
    import docx

    document = docx.Document()
    document.add_heading("Engineering Status Report - Q3 2026", level=1)

    document.add_paragraph(
        "Project Atlas (next-generation robotic arm) remains on track for a "
        "December beta release. Firmware integration testing is 70% complete."
    )
    document.add_paragraph(
        "Engineering headcount grew to 142 this quarter (+18), primarily in "
        "firmware and test engineering."
    )
    document.add_paragraph(
        "Unit test coverage improved to 81% across the firmware codebase, up "
        "from 74% last quarter."
    )
    document.add_heading("Supply Chain Risk", level=2)
    document.add_paragraph(
        "We've confirmed SensTech's proximity sensor lead-time increase "
        "mentioned in the finance report. The team has identified a second "
        "source, OptoSense, but full qualification and validation testing "
        "will not complete until early Q1 2027. Until then, Atlas arm "
        "shipments remain exposed to SensTech's lead times."
    )

    document.save(str(DOCS_DIR / "engineering_status.docx"))


def make_sales_pipeline_xlsx() -> None:
    import openpyxl

    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Pipeline"

    sheet.append(["Region", "Stage", "Deal Value (USD)", "Win Probability"])
    rows = [
        ["West", "Closed Won", 2100000, "100%"],
        ["East", "Negotiation", 4500000, "60%"],
        ["EU", "Discovery", 1200000, "20%"],
        ["APAC", "Closed Won", 900000, "100%"],
    ]
    for row in rows:
        sheet.append(row)

    sheet.append([])
    sheet.append(["Total pipeline value", "", 8700000, ""])
    sheet.append(["Weighted pipeline value", "", 6300000, ""])

    workbook.save(str(DOCS_DIR / "sales_pipeline.xlsx"))


def make_board_email_eml() -> None:
    msg = EmailMessage()
    msg["From"] = "ceo@acmerobotics.example"
    msg["To"] = "board@acmerobotics.example"
    msg["Subject"] = "Q3 Board Update - Customer Concentration & Retention"
    msg.set_content(
        "Board,\n\n"
        "A couple of items ahead of Thursday's meeting.\n\n"
        "Customer concentration: MetroLogistics remains our largest customer at "
        "24% of revenue. Their renewal conversation is underway and I'm cautiously "
        "optimistic, but I want the board aware of the concentration risk.\n\n"
        "Churn: overall customer churn ticked up to 4.1% this quarter, versus 2.8% "
        "last quarter. Two mid-market accounts cited pricing pressure from a new "
        "competitor, RoboFlex, as their reason for leaving.\n\n"
        "Proposal: I'd like to introduce a retention discount program for accounts "
        "above $500K ARR to get ahead of further RoboFlex-driven churn. Will bring "
        "a formal proposal Thursday.\n\n"
        "Best,\nCEO"
    )
    (DOCS_DIR / "board_email.eml").write_bytes(bytes(msg))


def make_allhands_notes_md() -> None:
    content = """# All-Hands Meeting Notes - September 2026

## Hiring
The company-wide hiring freeze remains in effect for Q4. **Exception approved:**
2 additional Robotics test engineer roles to support Project Atlas firmware
validation ahead of the December beta.

## People
Engineering had 3 resignations this quarter. Exit interviews consistently cited
compensation falling behind market rate as the primary reason for leaving.

## Facilities
The new espresso machine on the 3rd floor is finally working. Please report any
issues to facilities@acmerobotics.example.

## Reminder
Open enrollment for benefits closes October 15th.
"""
    (DOCS_DIR / "allhands_notes.md").write_text(content, encoding="utf-8")


def make_unrelated_file() -> None:
    content = """Office Snack Survey - Results

We asked the team what snacks they'd like restocked in the kitchen.

Top picks: trail mix (18 votes), sparkling water (15 votes), dark chocolate (12 votes).
Least popular: rice cakes (2 votes).

Thanks to everyone who voted! New snacks will be ordered next week.
"""
    (DOCS_DIR / "office_snack_survey.txt").write_text(content, encoding="utf-8")


def main() -> None:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    make_financial_summary_pdf()
    make_engineering_status_docx()
    make_sales_pipeline_xlsx()
    make_board_email_eml()
    make_allhands_notes_md()
    make_unrelated_file()
    print(f"Demo documents written to {DOCS_DIR}")
    for p in sorted(DOCS_DIR.iterdir()):
        print(f"  - {p.name}")


if __name__ == "__main__":
    main()
